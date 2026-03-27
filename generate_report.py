"""
Generate RiddleGame_MiniProjectReport.docx from JCOETSTORE_ReportPaper.docx template.
- Replaces text, team members, project-specific content
- DELETES all existing images/diagrams from the original docx
- Adds fresh AI-generated diagrams + game screenshots
"""

import docx
import os
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_PATH  = 'JCOETSTORE_ReportPaper.docx'
OUTPUT_PATH = 'RiddleGame_MiniProjectReport.docx'

# ─── Diagram paths (ai-generated, converted to proper PNG)
DIAGRAMS = [
    ('report_diagrams/sys_architecture_diagram_1774541190140_converted.png', 'Figure 1: System Architecture of Alphanumerical Riddle Challenge'),
    ('report_diagrams/user_workflow_diagram_1774541213974_converted.png',    'Figure 2: User Workflow Diagram'),
    ('report_diagrams/tech_stack_diagram_1774541231723_converted.png',       'Figure 3: Technology Stack Diagram'),
    ('report_diagrams/use_case_diagram_1774541249955_converted.png',         'Figure 4: Use Case Diagram'),
]

# ─── Game screenshots (converted to proper PNG)
SCREENSHOTS = [
    ('report_diagrams/screenshot_1.png', 'Figure 5: Welcome Page - Alphanumerical Riddle Challenge'),
    ('report_diagrams/screenshot_2.png', 'Figure 6: Stage Selection Page (Easy / Medium / Hard)'),
    ('report_diagrams/screenshot_3.png', 'Figure 7: Game Page - Active Riddle / Sequence'),
    ('report_diagrams/screenshot_4.png', 'Figure 8: Correct Answer Celebration'),
    ('report_diagrams/screenshot_5.png', 'Figure 9: Global Leaderboard'),
    ('report_diagrams/screenshot_6.png', 'Figure 10: Player Progress and Statistics Panel'),
    ('report_diagrams/screenshot_7.png', 'Figure 11: Level Map and Navigation'),
]

# ─── Text replacements (order matters — longer strings first to avoid partial matches)
REPLACEMENTS = [
    # Project Title
    ('JCOET STORE - E-COMMERCE PLATFORM', 'ALPHANUMERICAL RIDDLE CHALLENGE'),
    ('JCOET Store - E-Commerce Platform', 'Alphanumerical Riddle Challenge'),
    ('JCOET STORE - E-Commerce Platform', 'Alphanumerical Riddle Challenge'),
    ('JCOET Store', 'Alphanumerical Riddle Challenge'),
    ('JCOET STORE', 'ALPHANUMERICAL RIDDLE CHALLENGE'),
    ('E-Commerce Platform', 'Web-Based Interactive Game'),
    ('E-commerce Platform', 'Web-Based Interactive Game'),
    ('e-commerce platform', 'web-based interactive game'),
    ('e-Commerce Platform', 'Web-Based Interactive Game'),
    # Team Members (exact roll variations)
    ('Mr. Harshal Dhumane (24010054)', 'Ms. Shivani Akhare (25030024)'),
    ('Mr. Harshal Dhumane (24010053)', 'Ms. Shivani Akhare (25030024)'),
    ('Mr. Bhushan Angaitkar (24010338)', 'Mr. Sujal Bansode (25030031)'),
    ('Mr. Bhushan Angaikar (24010338)', 'Mr. Sujal Bansode (25030031)'),
    ('Mr. Atharav Dhole (24010062)', 'Mr. Vinayak Asole (25030042)'),
    ('Mr. Harshal Misale  (24010088)', 'Mr. Yash Awachat (25030036)'),
    ('Mr. Harshal Misale (24010088)', 'Mr. Yash Awachat (25030036)'),
    ('Mr. Daksh Dhawane (24010109)', ''),
    ('Mr. Daksh Dawane (24010109)', ''),
    # Numbered list names
    ('1. Mr. Harshal Dhumane (24010053)', '1. Ms. Shivani Akhare (25030024)'),
    ('2. Mr. Bhushan Angaikar (24010338)', '2. Mr. Sujal Bansode (25030031)'),
    ('3. Mr. Atharav Dhole (24010062)', '3. Mr. Vinayak Asole (25030042)'),
    ('4. Mr. Harshal Misale (24010088)', '4. Mr. Yash Awachat (25030036)'),
    ('5. Mr. Daksh Dhawane (24010109)', ''),
    # Bare names
    ('Mr. Harshal Dhumane', 'Ms. Shivani Akhare'),
    ('Mr. Bhushan Angaitkar', 'Mr. Sujal Bansode'),
    ('Mr. Bhushan Angaikar', 'Mr. Sujal Bansode'),
    ('Mr. Atharav Dhole', 'Mr. Vinayak Asole'),
    ('Mr. Harshal Misale', 'Mr. Yash Awachat'),
    ('Mr. Daksh Dhawane', ''),
    ('Mr. Daksh Dawane', ''),
    # Roll numbers standalone
    ('(24010053)', '(25030024)'),
    ('(24010054)', '(25030024)'),
    ('(24010338)', '(25030031)'),
    ('(24010062)', '(25030042)'),
    ('(24010088)', '(25030036)'),
    ('(24010109)', ''),
    # Project content
    ('shopping cart', 'riddle sequence'),
    ('Shopping Cart', 'Riddle Sequence'),
    ('Shopping Cart Management Screen', 'Stats & Progress Tracking Screen'),
    ('Shopping Cart Management', 'Progress Tracking'),
    ('Product Browsing and Search Interface', 'Riddle Selection Interface'),
    ('Product Details and Reviews', 'Riddle Details and Hints'),
    ('Checkout System', 'Score Submission & Level Completion'),
    ('User Profile', 'Player Profile'),
    ('Student Project Showcase', 'Game History and Rewards'),
    ('AI shopping assistant', 'Firebase Leaderboard'),
    ('Student Project Showcase', 'Level Achievement Showcase'),
    ('Loyalty points system', 'Streak reward system'),
    ('Recently viewed products', 'Recently solved riddles'),
    ('Responsive UI design', 'Animated responsive UI design'),
    ('campus e-commerce', 'educational gaming'),
    ('campus communities', 'student gamers'),
    ('campus life', 'educational challenges'),
    ('e-commerce solutions', 'gaming web-app solutions'),
    ('E-commerce', 'Gaming'),
    ('e-commerce', 'gaming'),
    ('wishlist', 'favorites list'),
    ('Wishlist', 'Favorites List'),
    ('inventory management', 'riddle database management'),
    ('Inventory management', 'Riddle database management'),
    ('inventory', 'riddle library'),
    ('Inventory', 'Riddle Library'),
    ('order tracking', 'score tracking'),
    ('Order tracking', 'Score tracking'),
    ('Product Data Structure', 'Riddle Data Structure'),
    ('product data', 'riddle data'),
    ('personalized homepages', 'personalized difficulty recommendations'),
    ('predictive inventory', 'adaptive difficulty'),
    ('chatbot customer support', 'in-game hint AI system'),
    ('social login', 'Firebase Google login'),
    ('product sharing', 'challenge sharing'),
    ('WebSocket connections for live inventory updates and order tracking',
     'Firebase Firestore real-time leaderboard and score sync'),
    ('Admin Dashboard', 'Admin Score Dashboard'),
    ('administrative interface for managing products, orders', 'admin interface for managing scores, users'),
    ('Backend Integration: Implement a full backend system with database for persistent data storage, real inventory management, and secure payment processing',
     'Backend Integration: Implement a full backend via Firebase for persistent leaderboard, real-time sync, and secure user authentication'),
    ('Mobile Application: Develop native mobile apps for iOS and Android platforms',
     'Mobile Application: Develop a Progressive Web App (PWA) for mobile devices'),
    ('Social Features: Add social login, product sharing, and community recommendations',
     'Social Features: Add challenge sharing, in-app chat, and competitive multiplayer mode'),
    ('Analytics Dashboard: Provide detailed analytics on sales, user behavior, and trends',
     'Analytics Dashboard: Provide detailed analytics on player performance, completion rates, and difficulty trends'),
    ('Internationalization: Add multi-language support for diverse campus communities',
     'Internationalization: Add multi-language riddles for diverse student communities'),
    # Sections renaming
    ('JCOET Store project represents a significant step forward in campus e-commerce solutions',
     'Alphanumerical Riddle Challenge represents a significant step forward in educational gaming web applications'),
    ('By combining traditional shopping functionality with modern AI capabilities and community features, it creates a platform that serves not just as a store, but as a digital hub for campus life.',
     'By combining classic logic puzzles with modern Firebase integration, real-time leaderboards, and smooth animations, it creates an engaging learning environment that challenges students cognitively.'),
    ('the project demonstrates that sophisticated e-commerce solutions can be built using accessible web technologies, making them viable for educational institutions of all sizes.',
     'The project demonstrates that sophisticated educational games can be built using only HTML, CSS, and JavaScript with Firebase, making them highly accessible and deployable for institutions of all sizes.'),
    # Gemini API ref
    ('Google, \"Gemini API Documentation,\" Google AI Studio, https://ai.google.dev/docs',
     'Google, \"Firebase Documentation,\" Firebase Docs, https://firebase.google.com/docs'),
    # Code appendix
    ('addToCart', 'checkAnswer'),
    ('updateCart', 'updateStats'),
    ('removeFromCart', 'skipLevel'),
    ('placeOrder', 'markLevelCompleted'),
    ('toggleWishlist', 'toggleHint'),
    ('addReview', 'recordScore'),
    ('callGeminiApi', 'syncToFirestore'),
    ("localStorage.setItem('cart', JSON.stringify(cart));",
     "localStorage.setItem('riddleGameStats', JSON.stringify(stats));"),
    # Misc project details
    ('campus', 'student'),
    ('product', 'riddle'),
    ('products', 'riddles'),
    ('Product', 'Riddle'),
    ('Products', 'Riddles'),
    ('orders', 'scores'),
    ('order', 'score'),
    ('Orders', 'Scores'),
    ('Order', 'Score'),
    ('store', 'game'),
    ('Store', 'Game'),
    ('cart', 'answer'),
    ('Cart', 'Answer'),
]

# ──────────────────────────────────────────────
# Helper: replace text in a run while preserving fmt
# ──────────────────────────────────────────────
def replace_in_para(paragraph):
    full_text = ''.join(r.text for r in paragraph.runs)
    changed = False
    for old, new in REPLACEMENTS:
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True
    if changed:
        # Rebuild runs: put everything in first run, clear rest
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for r in paragraph.runs[1:]:
                r.text = ''
    return changed


def process_paragraphs(paragraphs):
    for p in paragraphs:
        replace_in_para(p)


def process_doc(doc):
    # Body paragraphs
    process_paragraphs(doc.paragraphs)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
    # Headers & Footers
    for section in doc.sections:
        for hdr_or_ftr in [section.header, section.footer,
                           section.even_page_header, section.even_page_footer,
                           section.first_page_header, section.first_page_footer]:
            try:
                if hdr_or_ftr:
                    process_paragraphs(hdr_or_ftr.paragraphs)
            except Exception:
                pass


# ──────────────────────────────────────────────
# Helper: remove ALL images from docx
# ──────────────────────────────────────────────
def remove_all_images(doc):
    """Remove every inline image and drawing from the document body."""
    removed = 0
    # Walk all paragraphs and remove runs that contain pictures
    for para in doc.paragraphs:
        runs_to_remove = []
        for run in para.runs:
            # Check for drawing/picture XML inside run
            r_xml = run._r
            drawings = r_xml.findall('.//' + qn('w:drawing'))
            picts    = r_xml.findall('.//' + qn('w:pict'))
            if drawings or picts:
                runs_to_remove.append(r_xml)
                removed += 1
        for r_xml in runs_to_remove:
            r_xml.getparent().remove(r_xml)

    # Also strip inline shapes from table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    runs_to_remove = []
                    for run in para.runs:
                        r_xml = run._r
                        drawings = r_xml.findall('.//' + qn('w:drawing'))
                        picts    = r_xml.findall('.//' + qn('w:pict'))
                        if drawings or picts:
                            runs_to_remove.append(r_xml)
                            removed += 1
                    for r_xml in runs_to_remove:
                        r_xml.getparent().remove(r_xml)

    # Remove entire paragraphs that only had images (now empty) but keep structure
    print(f'  >> Removed {removed} image element(s) from document body.')


# ──────────────────────────────────────────────
# Helper: add a styled heading
# ──────────────────────────────────────────────
def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x8c)
    return heading


# ──────────────────────────────────────────────
# Helper: add image with caption
# ──────────────────────────────────────────────
def add_image_with_caption(doc, img_path, caption, width_inches=6.0):
    if not os.path.exists(img_path):
        print(f'  ⚠ Image not found: {img_path}')
        doc.add_paragraph(f'[Image not found: {os.path.basename(img_path)}]')
        return
    try:
        doc.add_picture(img_path, width=Inches(width_inches))
        # Center the picture
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Caption
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.runs[0] if cap.runs else cap.add_run(caption)
        cap_run.font.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        doc.add_paragraph('')  # spacing
    except Exception as e:
        print(f'  ⚠ Error adding {img_path}: {e}')
        doc.add_paragraph(f'[Error loading image: {e}]')


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────
def main():
    print(f'Loading template: {INPUT_PATH}')
    doc = Document(INPUT_PATH)

    print('Step 1: Removing all existing images/diagrams...')
    remove_all_images(doc)

    print('Step 2: Replacing text content...')
    process_doc(doc)

    # Apply font improvements to all body paragraphs
    print('Step 3: Enhancing paragraph fonts...')
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name in (None, '', 'Calibri', 'Times New Roman', 'Arial'):
                run.font.name = 'Calibri'
            if run.font.size is None:
                run.font.size = Pt(11)

    # ── Add PAGE BREAK and Diagrams Section
    print('Step 4: Adding fresh diagrams and screenshots...')
    doc.add_page_break()
    add_section_heading(doc, '10. DIAGRAMS AND SYSTEM ARCHITECTURE', level=1)

    for img_path, caption in DIAGRAMS:
        doc.add_paragraph('')
        add_image_with_caption(doc, img_path, caption, width_inches=6.0)

    # ── Add Screenshots Section
    doc.add_page_break()
    add_section_heading(doc, '11. GAME SCREENSHOTS', level=1)
    p = doc.add_paragraph(
        'The following screenshots illustrate the key interfaces and features of the '
        'Alphanumerical Riddle Challenge web application.'
    )
    p.runs[0].font.size = Pt(11)
    doc.add_paragraph('')

    for img_path, caption in SCREENSHOTS:
        add_image_with_caption(doc, img_path, caption, width_inches=5.5)

    print(f'Saving to: {OUTPUT_PATH}')
    doc.save(OUTPUT_PATH)
    print(f'\n✓ Done! Saved → {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
