import docx
import sys

def read_docx(file_path):
    doc = docx.Document(file_path)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P{i}: {p.text}")
    print("\n--- TABLES ---")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                print(f"T{t_idx}R{r_idx}: {' | '.join(row_data)}")

if __name__ == '__main__':
    read_docx(sys.argv[1])
