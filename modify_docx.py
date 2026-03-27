import docx
import os

def replace_text_in_paragraph(paragraph, replacements):
    text = paragraph.text
    original_text = text
    for old, new in replacements.items():
        if isinstance(old, str):
            text = text.replace(old, new)
            
    if text != original_text:
        # Save paragraph formatting
        if len(paragraph.runs) > 0:
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_bold = first_run.font.bold
            font_italic = first_run.font.italic
            
            paragraph.clear()
            new_run = paragraph.add_run(text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.font.bold = font_bold
            new_run.font.italic = font_italic
        else:
            paragraph.text = text

def process_blocks(blocks, replacements):
    for block in blocks:
        for para in block.paragraphs:
            replace_text_in_paragraph(para, replacements)
        for table in block.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_text_in_paragraph(para, replacements)

def modify_document(input_path, output_path):
    doc = docx.Document(input_path)
    
    replacements = {
        "JCOET STORE - E-COMMERCE PLATFORM": "ALPHANUMERICAL RIDDLE CHALLENGE",
        "JCOET Store - E-Commerce Platform": "Alphanumerical Riddle Challenge",
        "JCOET STORE": "ALPHANUMERICAL RIDDLE CHALLENGE",
        "JCOET Store": "Alphanumerical Riddle Challenge",
        "E-Commerce Platform": "Web-Based Logic Game",
        "E-commerce Platform": "Web-Based Logic Game",
        "Shopping Cart Management": "Real-Time Leaderboard",
        "shopping cart": "leaderboard",
        "E-commerce": "Gaming",
        "e-commerce": "gaming",
        "Store": "Game",
        "store": "game",
        "Product Browsing and Search Interface": "Riddle Selection Interface",
        "product browsing": "riddle selection",
        "Product Details and Reviews": "Game Status and Instructions",
        "Checkout System": "Level Completion and Scoring",
        "User Profile": "Player Profile",
        "Student Project Showcase": "Game History and Rewards",
        "Checkout": "Score Submission",
        "checkout": "score submission",
        "Inventory management": "Puzzle generation",
        "inventory management": "puzzle generation",
        "Product Data Structure": "Puzzle Data Structure",
        "categories": "difficulty levels",
        "product": "puzzle",
        "products": "puzzles",
        "Orders": "Scores",
        "orders": "scores",
        "Order": "Score",
        "order": "score",
        # names block:
        "Mr. Harshal Dhumane (24010054)": "Ms. Shivani Akhare (25030024)",
        "Mr. Harshal Dhumane (24010053)": "Ms. Shivani Akhare (25030024)",
        "1. Mr. Harshal Dhumane (24010053)": "1. Ms. Shivani Akhare (25030024)",
        
        "Mr. Bhushan Angaitkar (24010338)": "Mr. Sujal Bansode (25030031)",
        "Mr. Bhushan Angaikar (24010338)": "Mr. Sujal Bansode (25030031)",
        "2. Mr. Bhushan Angaikar (24010338)": "2. Mr. Sujal Bansode (25030031)",
        
        "Mr. Atharav Dhole (24010062)": "Mr. Vinayak Asole (25030042)",
        "3. Mr. Atharav Dhole (24010062)": "3. Mr. Vinayak Asole (25030042)",
        
        "Mr. Harshal Misale  (24010088)": "Mr. Yash Awachat (25030036)",
        "Mr. Harshal Misale (24010088)": "Mr. Yash Awachat (25030036)",
        "4. Mr. Harshal Misale (24010088)": "4. Mr. Yash Awachat (25030036)",
        
        "Mr. Daksh Dhawane (24010109)": "",
        "Mr. Daksh Dawane (24010109)": "",
        "5. Mr. Daksh Dhawane (24010109)": "",
        
        "Mr. Harshal Dhumane": "Ms. Shivani Akhare",
        "Mr. Bhushan Angaitkar": "Mr. Sujal Bansode",
        "Mr. Bhushan Angaikar": "Mr. Sujal Bansode",
        "Mr. Atharav Dhole": "Mr. Vinayak Asole",
        "Mr. Harshal Misale": "Mr. Yash Awachat",
        "Mr. Daksh Dhawane": "",
        "Mr. Daksh Dawane": "",
        "(24010053)": "(25030024)",
        "(24010054)": "(25030024)",
        "(24010338)": "(25030031)",
        "(24010062)": "(25030042)",
        "(24010088)": "(25030036)",
        "(24010109)": "",
    }
    
    process_blocks([doc], replacements)
    
    # Process headers and footers
    for section in doc.sections:
        if section.header:
            process_blocks([section.header], replacements)
        if section.footer:
            process_blocks([section.footer], replacements)

    # Add screenshots section only if not already added
    if not any("SCREENSHOTS AND DIAGRAMS" in p.text for p in doc.paragraphs):
        doc.add_page_break()
        doc.add_heading('10. SCREENSHOTS AND DIAGRAMS', level=1)
        
        image_dir = 'images'
        if os.path.exists(image_dir):
            images = [f for f in os.listdir(image_dir) if f.endswith('.png') or f.endswith('.jpg')]
            images.sort()
            for img_name in images:
                img_path = os.path.join(image_dir, img_name)
                doc.add_paragraph(f"Figure: {img_name}")
                try:
                    from docx.shared import Inches
                    doc.add_picture(img_path, width=Inches(6.0))
                except Exception as e:
                    doc.add_paragraph(f"Error adding image: {e}")
                
    doc.save(output_path)
    print(f"Saved modified document to {output_path}")

if __name__ == '__main__':
    modify_document('JCOETSTORE_ReportPaper.docx', 'RiddleGame_ReportPaper.docx')
